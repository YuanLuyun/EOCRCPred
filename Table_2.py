import pandas as pd
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.enum.section import WD_ORIENTATION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, ns

# 加载 c-index95.csv 数据
file_path = 'c-index95.csv'  # 请替换为您的实际文件路径
c_index_data = pd.read_csv(file_path)

# 定义模型名称和 C-Index 列标签
model_names = ['RSF', 'CoxPH', 'S-SVM', 'XGBSE', 'GBSA', 'DeepSurv']
metrics = {
    'Train': ['Train C-index', 'Train C-index at 12 months', 'Train C-index at 36 months',
              'Train C-index at 60 months'],
    'Test': ['Test C-index', 'Test C-index at 12 months', 'Test C-index at 36 months', 'Test C-index at 60 months'],
    'External': ['External C-index', 'External C-index at 12 months', 'External C-index at 36 months',
                 'External C-index at 60 months']
}

# 创建 Word 文档
doc = Document()

# 设置页面为横向
section = doc.sections[-1]
section.orientation = WD_ORIENTATION.LANDSCAPE
section.page_width, section.page_height = section.page_height, section.page_width

# 添加标题
title_paragraph = doc.add_paragraph()
run = title_paragraph.add_run('Table 2. ')
run.bold = True  # 加粗 Table 2
run.font.name = 'Times New Roman'
run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
run = title_paragraph.add_run('C-index values at 12, 36, and 60 months across different models in three cohorts.')
run.font.name = 'Times New Roman'
run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

# 创建表格，行数为模型数量 + 1（标题行），列数为5列（模型名称，总体C-index，12个月C-index，36个月C-index，60个月C-index）
num_rows = len(model_names) * 3 + 1  # 每个模型三个数据集（训练、测试、外部验证）+ 标题行
num_cols = 5
table = doc.add_table(rows=1, cols=num_cols)

# 调整列宽
table.columns[0].width = Inches(2)  # 模型名称列
for i in range(1, num_cols):
    table.columns[i].width = Inches(2.1)  # C-index 列宽度


# 设置表格中所有单元格的字体为 Times New Roman
def set_table_font(table):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(11)
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')


# 创建标题行
header_cells = table.rows[0].cells
header_cells[0].text = 'Model'
header_cells[1].text = 'C-Index (95% CI)'
header_cells[2].text = '12 Months (95% CI)'
header_cells[3].text = '36 Months (95% CI)'
header_cells[4].text = '60 Months (95% CI)'

# 设置标题行格式
for i in range(num_cols):
    paragraph = header_cells[i].paragraphs[0]
    run = paragraph.runs[0]
    run.bold = True
    if i == 0:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT  # Left-align the first column
    else:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Center-align other columns


# 按顺序先填充训练集、再填充测试集、最后填充外部验证集
for dataset in ['Train', 'Test', 'External']:
    # 在数据集开头添加一行表示数据集名称
    row_cells = table.add_row().cells
    if dataset == 'Train':
        row_cells[0].text = 'Training cohort'
    elif dataset == 'Test':
        row_cells[0].text = 'Test cohort'
    else:
        row_cells[0].text = 'External validation cohort'

    # 合并该行的所有单元格
    for i in range(1, num_cols):
        row_cells[i].merge(row_cells[0])

    # 设置加粗和居左
    paragraph = row_cells[0].paragraphs[0]
    run = paragraph.runs[0]  # 使用已有的文本内容
    run.bold = True
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    for model in model_names:
        row_cells = table.add_row().cells
        row_cells[0].text = f'{model}'

        # 设置“Model”列居左
        paragraph = row_cells[0].paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # 确保文本的 run 也与居左对齐
        run = paragraph.runs[0]
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)

        # 继续填充其他列
        for i, metric in enumerate(metrics[dataset], start=1):
            value = c_index_data.loc[c_index_data['Unnamed: 0'] == metric, model].values[0]
            row_cells[i].text = value

        # 其他列居中对齐
        for i in range(1, num_cols):
            row_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

# 设置表格字体
set_table_font(table)


# 设置表头下方添加底部边框（1磅 = 8）
def set_header_bottom_border(table):
    for cell in table.rows[0].cells:
        cell_borders = OxmlElement('w:tcBorders')
        bottom = OxmlElement('w:bottom')
        bottom.set(ns.qn('w:val'), 'single')
        bottom.set(ns.qn('w:sz'), '8')  # 1磅，尺寸为8
        bottom.set(ns.qn('w:color'), '000000')  # 边框颜色为黑色
        cell_borders.append(bottom)
        cell._tc.get_or_add_tcPr().append(cell_borders)


# 应用表头底部边框
set_header_bottom_border(table)


# 设置整个表格的顶部和底部边框（1.5磅 = 12）
def set_table_top_bottom_border(table):
    tbl = table._tbl
    tblBorders = OxmlElement('w:tblBorders')
    top_border = OxmlElement('w:top')
    top_border.set(ns.qn('w:val'), 'single')
    top_border.set(ns.qn('w:sz'), '12')  # 1.5磅，尺寸为12
    top_border.set(ns.qn('w:color'), '000000')
    tblBorders.append(top_border)
    bottom_border = OxmlElement('w:bottom')
    bottom_border.set(ns.qn('w:val'), 'single')
    bottom_border.set(ns.qn('w:sz'), '12')  # 1.5磅，尺寸为12
    bottom_border.set(ns.qn('w:color'), '000000')
    tblBorders.append(bottom_border)
    tbl.tblPr.append(tblBorders)


# 应用表格顶部和底部边框
set_table_top_bottom_border(table)

# 保存 Word 文档
doc.save('Table 2.docx')

print("Table 2.docx")
