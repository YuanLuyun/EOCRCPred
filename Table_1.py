import pandas as pd
from docx import Document
from docx.shared import Pt  # 用于设置字体大小
from docx.oxml.ns import qn  # 用于设置字体
from scipy.stats import chi2_contingency, fisher_exact
from docx.enum.section import WD_ORIENTATION
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# 加载 external_validation_set.csv 数据并去掉 'Patient_ID' 列
external_data = pd.read_csv('external_validation_set_unencoded.csv')
external_data.columns = external_data.columns.str.replace('_', ' ')

# 加载 data_encoded7408.csv 数据并去掉 'Patient_ID' 列
data = pd.read_csv('data_unencoded7408.csv')
data.columns = data.columns.str.replace('_', ' ')
data = data.drop(columns=['Patient ID'])  # 删除 Patient_ID 列

# 特征与标签的分离，排除 'Survival status' 和 'OS month'
X = data.drop(columns=['Survival status', 'OS month'])  # 假设 'Survival status' 是标签，'OS month' 不纳入统计
y = data['Survival status']

# 按照 9:1 划分训练集与测试集
from sklearn.model_selection import train_test_split
X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.2, random_state=42)

# 创建 Word 文档
doc = Document()
# 设置页面为横向
section = doc.sections[-1]
section.orientation = WD_ORIENTATION.LANDSCAPE

# 设置页面宽度和高度为横向的值
new_width, new_height = section.page_height, section.page_width
section.page_width = new_width
section.page_height = new_height
# 添加标题，字体与表格一致
title_paragraph = doc.add_paragraph()
run = title_paragraph.add_run('Table 1. ')
run.bold = True  # 使 Table 1 加粗
run.font.name = 'Times New Roman'  # 设置字体为 Times New Roman
run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')  # 确保所有语言的字体都是 Times New Roman
run = title_paragraph.add_run('Baseline clinicopathological characteristics of the training, Testing, and external Validation cohort.')
run.font.name = 'Times New Roman'
run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

# 创建表格，行数会根据特征和分类特征的类别数量动态增加
table = doc.add_table(rows=1, cols=7)  # 修改列数为7
# 调整列宽，插入的"Overall"列宽度设置与其他列相同
table.columns[0].width = Inches(2.5)  # 特征列宽度不变
table.columns[1].width = Inches(1.6)  # Overall cohort 新增的列
table.columns[2].width = Inches(1.6)  # Training cohort
table.columns[3].width = Inches(1.6)  # Testing cohort
table.columns[4].width = Inches(1.5)  # External Validation cohort
table.columns[5].width = Inches(1.4)  # p-value (Training vs Testing)
table.columns[6].width = Inches(1.4)  # p-value (Training vs External)

# 设置表格中所有单元格的字体为 Times New Roman
def set_table_font(table):
    """
    设置表格中所有单元格的字体为 Times New Roman
    """
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(11)
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')  # 设置东亚字体


first_row_cells = table.rows[0].cells

# 合并第2到第4个单元格为 "Seer cohort"
first_row_cells[1].merge(first_row_cells[3])
p = first_row_cells[1].paragraphs[0]
run = p.add_run("SEER cohort\nN(%)")
run.bold = True
p.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 居中对齐

# 在第6个单元格插入 "External Validation cohort"
p = first_row_cells[4].paragraphs[0]
run = p.add_run("External validation cohort N(%)")
run.bold = True
p.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 居中对齐
# 在第二行插入具体表头之前，先添加一行
second_row = table.add_row()  # 添加第二行
# 第二行插入具体表头，并在 Testing cohort 和 p-value 之间留空白列
second_row_cells = table.rows[1].cells
headers = [
    'Characteristics',
    'Overall',
    'Training cohort',
    'Test cohort',
    '',  # 空白列
    'p-value\n(Training vs Test)',
    'p-value\n(Training vs External)'
]

# 填充第二行表头
for i, header in enumerate(headers):
    p = second_row_cells[i].paragraphs[0]
    run = p.add_run(header)
    run.bold = True
    if i > 0:  # 除了第一列，其余列居中对齐
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT  # 第一列左对齐
# 在表头下方添加一行
patients_row = table.add_row()  # 添加新行

# 计算样本数
total_samples = len(data)  # 总体样本数
train_samples = len(X_train)  # 训练集样本数
test_samples = len(X_test)  # 测试集样本数
external_samples = len(external_data)  # 外部验证集样本数

# 填充新行的内容
patients_cells = patients_row.cells
patients_cells[0].text = "Patients"  # 第一列写入 "Patients"
patients_cells[1].text = f"{total_samples}"  # 总体样本数
patients_cells[2].text = f"{train_samples}"  # 训练集样本数
patients_cells[3].text = f"{test_samples}"  # 测试集样本数
patients_cells[4].text = f"{external_samples}"  # 外部验证集样本数
patients_cells[5].text = ""  # p-value (Training vs Testing)，不适用于此行，留空
patients_cells[6].text = ""  # p-value (Training vs External)，不适用于此行，留空

# 使新行的文本居中
for i in range(1, 5):  # 从第二列到第四列
    patients_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

# 定义特征及其类别的顺序
feature_categories = {
    'Age': ['less than 35', '35-44', '45-49'],  # 保持数据集中原始分类
    'Sex': ['Male', 'Female'],
    'Race': ['White', 'Black', 'Other'],
    'Primary site': ['Ascending colon', 'Transverse colon', 'Descending colon', 'Sigmoid colon', 'Rectum'],
    'Tumor size': ['less than 5', '5+'],
    'Grade': ['Well differentiated', 'Moderately differentiated', 'Poorly differentiated', 'Undifferentiated'],
    'Histology': ['Non-specific adenocarcinoma', 'Specific adenocarcinoma', 'Other'],
    'TNM Stage': ['0', 'I', 'IIA', 'IIB', 'IIC', 'IIIA', 'IIIB', 'IIIC'],
    'T': ['Tis', 'T1', 'T2', 'T3', 'T4'],
    'N': ['N0', 'N1', 'N2'],
    'CEA': ['negative', 'Borderline', 'positive'],
    'Resection type': ['Partial/subtotal colectomy', 'Hemicolectomy or greater', 'Total colectomy', 'Colectomy plus removal of other organs'],
    'No.of resected LNs': ['Zero', '1 to 3', '4+'],
    'Tumor Deposits': ['Zero', '1 to 2', '3+'],
    'Perineural Invasion': ['No', 'Yes'],
    'Surg.Rad.Seq': ['Untreated', 'Postoperative', 'Preoperative', 'Preoperative+Postoperative', 'Sequence unknown'],
    'Chemotherapy': ['No/Unknown', 'Yes'],
    'Systemic.Sur.Seq': ['Untreated', 'Postoperative', 'Preoperative', 'Preoperative+Postoperative', 'Sequence unknown'],
    'Marital status': ['Single', 'Married', 'Divorced', 'Widowed'],
    'Median household income': ['less than $35,000', '$35,000-$54,999', '$55,000-$74,999', '$75,000+']
}
replace_dict = {
    'Age': 'Age (years)',
    'Tumor size': 'Tumor size (cm)',
    'less than 35':'＜35',
    'less than 5':'＜5',
    '5+':'≥5',
    'Zero':'0',
    '1 to 3':'1-3',
    '4+':'≥4',
    '1 to 2':'1-2',
    '3+':'≥3',
    'less than $35,000':'＜$35,000',
    '$75,000+':'≥$75,000',
    'CEA':'CEA(ng/ml)',
    'negative':'＜5',
    'Borderline':'5',
    'positive':'＞5',
}
# 创建一个 DataFrame 来存储每个类别的样本数
data_list = []
# 分类和数值特征的分类统计
for feature, categories in feature_categories.items():
    # 初始化列联表
    contingency_table_train_test = []
    contingency_table_train_external = []
    # 只在表格显示时使用替换后的名称，但在数据处理中使用原始名称
    feature_display_name = replace_dict.get(feature, feature)  # 表格显示时使用替换后的名称
    # 添加特征名称行
    row_cells = table.add_row().cells
    p = row_cells[0].paragraphs[0]
    run = p.add_run(feature_display_name)  # 在表格中显示替换后的特征名称

    # 针对每个类别进行样本数统计并显示 n(%)
    for category in categories:
        # 在显示类别时应用替换
        category_display_name = replace_dict.get(category, category)  # 表格显示时使用替换后的类别名称
        row_cells = table.add_row().cells
        p = row_cells[0].paragraphs[0]  # 获取第一个段落
        run = p.add_run(category_display_name)  # 使用替换后的类别名称

        # 设置左缩进，增加缩进量
        p.paragraph_format.left_indent = Inches(0.2)  # 左缩进设置为 0.2 英寸

        # 计算训练集样本数和百分比
        train_count = (X_train[feature] == category).sum()
        train_percent = train_count / len(X_train) * 100
        row_cells[2].text = f"{train_count} ({train_percent:.2f}%)"  # 显示 n(%)

        # 计算测试集样本数和百分比
        test_count = (X_test[feature] == category).sum()
        test_percent = test_count / len(X_test) * 100
        row_cells[3].text = f"{test_count} ({test_percent:.2f}%)"  # 显示 n(%)
        # Overall cohort (训练集+测试集)
        overall_count = train_count + test_count
        overall_percent = overall_count / (len(X_train) + len(X_test)) * 100
        row_cells[1].text = f"{overall_count} ({overall_percent:.2f}%)"  # 第2列显示总体 n(%)
        # 计算外部验证集样本数和百分比
        external_count = (external_data[feature] == category).sum()
        external_percent = external_count / len(external_data) * 100
        row_cells[4].text = f"{external_count} ({external_percent:.2f}%)"  # 显示 n(%)

        # 将结果添加到数据列表中
        data_list.append([feature, category, train_count, test_count, external_count])

        # 将样本数添加到列联表中
        contingency_table_train_test.append([train_count, test_count])
        contingency_table_train_external.append([train_count, external_count])
    # 打印列联表
    print(f"\nContingency Table (Train vs Test) for feature: {feature}")
    print(contingency_table_train_test)
    print(f"\nContingency Table (Train vs External) for feature: {feature}")
    print(contingency_table_train_external)

    # 对当前特征执行卡方检验 (训练集 vs 测试集)
    try:
        chi2_train_test, p_value_train_test, _, _ = chi2_contingency(contingency_table_train_test)
        print(f"Feature: {feature} (Train vs Test), Chi2: {chi2_train_test:.3f}, P-value: {p_value_train_test:.3f}")
    except ValueError:
        print(f"Feature: {feature} (Train vs Test), 无法执行卡方检验 (可能数据为0或单一类别)")

    # 对当前特征执行卡方检验 (训练集 vs 外部验证集)
    try:
        chi2_train_external, p_value_train_external, _, _ = chi2_contingency(contingency_table_train_external)
        print(f"Feature: {feature} (Train vs External), Chi2: {chi2_train_external:.3f}, P-value: {p_value_train_external:.3f}")
    except ValueError:
        print(f"Feature: {feature} (Train vs External), 无法执行卡方检验 (可能数据为0或单一类别)")
    # 将 p-value 写入 Word 文档表格的第五列和第六列
    row_cells[5].text = f"{p_value_train_test:.3f}" if isinstance(p_value_train_test, float) else p_value_train_test
    row_cells[6].text = f"{p_value_train_external:.3f}" if isinstance(p_value_train_external,
                                                                          float) else p_value_train_external

# 填充完表格后，设置第一列以外的内容居中对齐
for row in table.rows:
    for i, cell in enumerate(row.cells):
        for paragraph in cell.paragraphs:
            if i == 0:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT  # 第一列左对齐
            else:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 其余列居中对齐
def set_table_border(table):
    """
    设置表格的边框：仅保留整个表格的顶部、表头下方、以及表格最底部的边框。
    表格最上和最下的边框为1.5磅，表头下方的边框为1磅。
    """
    tbl = table._tbl  # 获取表格的 XML 元素
    tblBorders = OxmlElement('w:tblBorders')

    # 设置表格最上方的边框（1.5磅 = 12）
    top_border = OxmlElement('w:top')
    top_border.set(ns.qn('w:val'), 'single')
    top_border.set(ns.qn('w:sz'), '12')  # 1.5磅，尺寸为12
    top_border.set(ns.qn('w:color'), '000000')  # 边框颜色为黑色
    tblBorders.append(top_border)

    # 设置表格最下方的边框（1.5磅 = 12）
    bottom_border = OxmlElement('w:bottom')
    bottom_border.set(ns.qn('w:val'), 'single')
    bottom_border.set(ns.qn('w:sz'), '12')  # 1.5磅，尺寸为12
    bottom_border.set(ns.qn('w:color'), '000000')  # 边框颜色为黑色
    tblBorders.append(bottom_border)

    # 应用边框样式
    tbl.tblPr.append(tblBorders)

from docx.oxml import OxmlElement, ns

# 为表头第二行的每个单元格添加底部边框（1磅 = 8）
for cell in table.rows[1].cells:
    cell_borders = OxmlElement('w:tcBorders')

    # 设置单元格底部的边框（1磅 = 8）
    bottom = OxmlElement('w:bottom')
    bottom.set(ns.qn('w:val'), 'single')
    bottom.set(ns.qn('w:sz'), '8')  # 1磅，尺寸为8
    bottom.set(ns.qn('w:color'), '000000')  # 边框颜色为黑色
    cell_borders.append(bottom)

    # 应用到单元格
    cell._tc.get_or_add_tcPr().append(cell_borders)

# 为第一行的第2、3、4列添加底部边框
for i in range(1, 4):  # 第一行的第2到第4列 (索引1到3)
    cell_borders = OxmlElement('w:tcBorders')

    # 设置单元格底部的边框（1磅 = 8）
    bottom = OxmlElement('w:bottom')
    bottom.set(ns.qn('w:val'), 'single')
    bottom.set(ns.qn('w:sz'), '8')  # 1磅，尺寸为8
    bottom.set(ns.qn('w:color'), '000000')  # 边框颜色为黑色
    cell_borders.append(bottom)

    # 应用到单元格
    table.rows[0].cells[i]._tc.get_or_add_tcPr().append(cell_borders)


# 应用三线表样式
set_table_border(table)
# 调用该函数来设置整个表格的字体
set_table_font(table)
# 保存表格之前，在表格下方添加注释
paragraph = doc.add_paragraph()
run = paragraph.add_run('Abbreviations:')
run.bold = True  # 加粗 "Abbreviations"
run.font.name = 'Times New Roman'  # 设置字体为 Times New Roman
run.font.size = Pt(11)
run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

# 添加缩写解释
abbreviation_text = """
CEA: Carcinoembryonic Antigen; No. of resected LNs: Number of Resected Lymph Nodes; Surg.Rad.Seq: Surgery Radiation Sequence; Systemic.Sur.Seq: Systemic Surgery Sequence
"""

paragraph.add_run(abbreviation_text).font.size = Pt(11)

# 设置注释字体和格式
for p in doc.paragraphs:
    if p.text.startswith('Abbreviations:'):
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

# 保存 Word 文档
doc.save('Table 1.docx')

print("三线表已生成，并保存为 Table 1.docx")
