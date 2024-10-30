import pandas as pd
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.enum.section import WD_ORIENTATION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, ns

# 加载 Time-dependent AUC 数据
file_path = 'Time dependent AUC.csv'
auc_data = pd.read_csv(file_path)

# 定义模型名称
model_names = ['RSF', 'CoxPH', 'S-SVM', 'XGBSE', 'GBSA', 'DeepSurv']

# 定义数据集名称
datasets = ['Training cohort', 'Test cohort', 'External validation cohort']

# 手动定义列顺序，无论数据中的时间点，按照这个顺序来填充表格列
column_order = [6, 12, 18, 24, 30, 36, 42, 48, 54, 60, 66, 72, 78, 84, 90, 96, 102, 108, 114, 120, 'Mean AUC']

# 创建 Word 文档
doc = Document()

# 设置页面为横向
section = doc.sections[-1]
section.orientation = WD_ORIENTATION.LANDSCAPE
section.page_width, section.page_height = section.page_height, section.page_width

# 添加标题
title_paragraph = doc.add_paragraph()
run = title_paragraph.add_run('Table 3. ')
run.bold = True  # 加粗
run.font.size = Pt(8)  # 设置标题字体大小
run.font.name = 'Times New Roman'
run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
run = title_paragraph.add_run('Time-dependent AUC values across different models in three cohorts.')
run.font.size = Pt(8)  # 将字体大小设置为 12pt
run.font.name = 'Times New Roman'
run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

# 创建表格，行数为模型数量 * 3（3个数据集） + 1（标题行），列数为指定的时间点顺序
num_rows = 2  # 仅包括两行标题行
num_cols = len(column_order) + 1  # 列数为指定的时间点 + 模型列
table = doc.add_table(rows=num_rows, cols=num_cols)

# 关闭表格的自动调整功能
table.autofit = False
# 调整列宽
table.columns[0].width = Inches(0.6)  # 模型名称列
table.columns[-1].width = Inches(0.5)  # 最后一列 'Mean AUC'
for i in range(1, num_cols - 1):
    table.columns[i].width = Inches(0.4)  # AUC 列宽度

# 设置表格中所有单元格的字体为 Times New Roman
def set_table_font(table):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(8)
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

# 获取第一行
top_row = table.rows[0].cells

# 合并第二个到倒数第二个单元格，并填入 'AUC (month)'
for i in range(1, num_cols - 1):
    top_row[i].merge(top_row[1])  # 将第i个单元格合并到第2个单元格

# 将合并后的第二个单元格内容设置为 'AUC (month)'
top_row[1].text = 'AUC (month)'

# 设置第一列的内容
top_row[0].text = "Model"

# 获取第二行作为新的表头行
header_cells = table.rows[1].cells

# 填入其他列的表头，按照指定的顺序
for i, time in enumerate(column_order):
    if time == 'Mean AUC':
        header_cells[i + 1].text = 'mean'  # 将最后一列的表头设置为 'Mean AUC'
    else:
        header_cells[i + 1].text = str(time)  # 其他列按顺序填入时间点


# 定义一个统一的格式设置函数
def set_row_format(cells):
    for i in range(num_cols):
        paragraph = cells[i].paragraphs[0]  # 获取当前单元格的段落
        if paragraph.runs:  # 检查段落是否有内容
            run = paragraph.runs[0]  # 获取段落中的运行内容
            run.bold = True  # 将文本加粗

        # 设置对齐方式，第一列左对齐，其他列居中对齐
        if i == 0:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT  # 左对齐第一列
        else:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 其他列居中对齐

# 统一设置第一行和第二行的格式
set_row_format(top_row)  # 设置第一行的格式
set_row_format(header_cells)  # 设置第二行的格式

# 填充数据
for dataset in datasets:
    # 在数据集开头添加一行表示数据集名称
    row_cells = table.add_row().cells
    row_cells[0].text = dataset

    # 合并该行的所有单元格
    for i in range(1, num_cols):
        row_cells[i].merge(row_cells[0])

    # 设置加粗和居左
    paragraph = row_cells[0].paragraphs[0]
    run = paragraph.runs[0]  # 使用已有的文本内容
    run.bold = True
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # 填充模型数据
    for model in model_names:
        row_cells = table.add_row().cells
        row_cells[0].text = f'{model}'

        # 设置“Model”列居左
        paragraph = row_cells[0].paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

        for i, time in enumerate(column_order[:-1], start=1):
            # 将time转换为字符串
            time_str = str(time)
            # 提取对应时间点的AUC分数
            result = auc_data.loc[
                (auc_data['Dataset'] == dataset) & (auc_data['Model'] == model) & (auc_data['Times'] == time_str)
                ]
            if not result.empty:
                auc_value = result['AUC Scores'].values[0]
            else:
                auc_value = 0
            row_cells[i].text = f'{auc_value:.3f}'

        # 先将 'Times' 列转换为字符串类型，以便能够使用 .str.contains()
        auc_data['Times'] = auc_data['Times'].astype(str)
        # 直接从数据中提取 Mean AUC
        mean_auc_result = auc_data.loc[
            (auc_data['Dataset'] == dataset) & (auc_data['Model'] == model) & (auc_data['Times'].str.contains('mean', case=False, na=False))
        ]
        if not mean_auc_result.empty:
            mean_auc_value = mean_auc_result['AUC Scores'].values[0]
            row_cells[-1].text = f'{mean_auc_value:.3f}'
        else:
            row_cells[-1].text = 'N/A'

        # 其他列居中对齐
        for i in range(1, num_cols):
            row_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

# 设置表格字体
set_table_font(table)

# 设置表头下方添加底部边框
def set_header_bottom_border(table):
    # 遍历表格的前两行
    for row_index, row in enumerate(table.rows[:2]):
        for cell_index, cell in enumerate(row.cells):
            # 跳过第一行第一个单元格
            if row_index == 0 and cell_index == 0:
                continue
            cell_borders = OxmlElement('w:tcBorders')
            bottom = OxmlElement('w:bottom')
            bottom.set(ns.qn('w:val'), 'single')
            bottom.set(ns.qn('w:sz'), '8')  # 1磅，尺寸为8
            bottom.set(ns.qn('w:color'), '000000')  # 边框颜色为黑色
            cell_borders.append(bottom)
            cell._tc.get_or_add_tcPr().append(cell_borders)

# 应用表头底部边框
set_header_bottom_border(table)

# 设置整个表格的顶部和底部边框
def set_table_top_bottom_border(table):
    tbl = table._tbl
    tblBorders = OxmlElement('w:tblBorders')
    top_border = OxmlElement('w:top')
    top_border.set(ns.qn('w:val'), 'single')
    top_border.set(ns.qn('w:sz'), '12')  # 1.5磅，尺寸为12
    top_border.set(ns.qn('w:color'), '000000')
    tblBorders.append(top_border)
    bottom_border = OxmlElement('w:bottom')
    bottom_border.set(ns.qn('w:val'), 'single')
    bottom_border.set(ns.qn('w:sz'), '12')  # 1.5磅，尺寸为12
    bottom_border.set(ns.qn('w:color'), '000000')
    tblBorders.append(bottom_border)
    tbl.tblPr.append(tblBorders)

# 应用表格顶部和底部边框
set_table_top_bottom_border(table)

# 保存 Word 文档
doc.save('Table 3.docx')