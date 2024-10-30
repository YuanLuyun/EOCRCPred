from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.enum.section import WD_ORIENTATION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, ns
import pandas as pd
from docx.shared import Mm
# 读取RSF特征重要性数据
file_path = 'RSF_permutation_importance.csv'
rsf_data = pd.read_csv(file_path)
# 替换所有数据中的下划线
rsf_data = rsf_data.applymap(lambda x: x.replace('_', ' ') if isinstance(x, str) else x)

# 创建Word文档
doc = Document()

# 设置纵向页面
section = doc.sections[-1]
section.orientation = WD_ORIENTATION.PORTRAIT
# 将页面尺寸设置为标准的A4纵向尺寸
section.page_width = Mm(210)   # A4宽度为210mm
section.page_height = Mm(297)  # A4高度为297mm
# 添加标题
title_paragraph = doc.add_paragraph()
run = title_paragraph.add_run('Supplementary Table 1. ')
run.bold = True
run.font.name = 'Times New Roman'
run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
run = title_paragraph.add_run('Relative importance of variables in the RSF model based on permutation importance scores.')
run.font.name = 'Times New Roman'
run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

# 创建表格：2 列（特征名称，特征重要性）
num_rows = len(rsf_data) + 1  # 包含表头
num_cols = 2
table = doc.add_table(rows=1, cols=num_cols)

# 调整列宽
table.columns[0].width = Inches(3)  # 特征名称列
table.columns[1].width = Inches(2)  # 特征重要性列

# 设置表格字体
def set_table_font(table):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(11)
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

# 设置三线表的边框
def set_table_border(table):
    # 设置表头上方的线（第一条线）
    tbl = table._tbl
    top_border = OxmlElement('w:top')
    top_border.set(ns.qn('w:val'), 'single')
    top_border.set(ns.qn('w:sz'), '12')  # 1.5pt, size 12
    top_border.set(ns.qn('w:color'), '000000')
    tblBorders = OxmlElement('w:tblBorders')
    tblBorders.append(top_border)

    # 设置表格底部的线（第三条线）
    bottom_border = OxmlElement('w:bottom')
    bottom_border.set(ns.qn('w:val'), 'single')
    bottom_border.set(ns.qn('w:sz'), '12')  # 1.5pt, size 12
    bottom_border.set(ns.qn('w:color'), '000000')
    tblBorders.append(bottom_border)
    tbl.tblPr.append(tblBorders)

# 设置表头底部的边框（第二条线）
def set_header_bottom_border(table):
    header_cells = table.rows[0].cells
    for cell in header_cells:
        tcBorders = OxmlElement('w:tcBorders')
        bottom_border = OxmlElement('w:bottom')
        bottom_border.set(ns.qn('w:val'), 'single')
        bottom_border.set(ns.qn('w:sz'), '8')  # 1pt, size 8
        bottom_border.set(ns.qn('w:color'), '000000')
        tcBorders.append(bottom_border)
        cell._tc.get_or_add_tcPr().append(tcBorders)

# 创建表头
header_cells = table.rows[0].cells
header_cells[0].text = 'Feature'
header_cells[1].text = 'Permutation importance'

# 设置表头字体加粗，并且调整对齐方式
# 第一列左对齐
header_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
header_cells[0].paragraphs[0].runs[0].bold = True

# 第二列居中对齐
header_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
header_cells[1].paragraphs[0].runs[0].bold = True

# 填充数据
for index, row in rsf_data.iterrows():
    row_cells = table.add_row().cells
    row_cells[0].text = row['Feature']
    row_cells[1].text = f"{row['Permutation importance']:.6f}"  # 格式化显示特征重要性

    # 设置每列的对齐方式
    row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT  # 左对齐特征名称
    row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER  # 居中对齐特征重要性

# 设置表格字体
set_table_font(table)

# 应用三线表边框
set_table_border(table)
set_header_bottom_border(table)  # 设置表头底部边框
# 保存表格之前，在表格下方添加注释
paragraph = doc.add_paragraph()
run = paragraph.add_run('Abbreviations:')
run.bold = True  # 加粗 "Abbreviations"
run.font.name = 'Times New Roman'  # 设置字体为 Times New Roman
run.font.size = Pt(11)
run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

# 添加缩写解释
abbreviation_text = """
CEA: Carcinoembryonic Antigen; No. of resected LNs: Number of Resected Lymph Nodes; OS: Overall Survival; Surg.Rad.Seq: Surgery Radiation Sequence; Systemic.Sur.Seq: Systemic Surgery Sequence
"""

paragraph.add_run(abbreviation_text).font.size = Pt(11)

# 设置注释字体和格式
for p in doc.paragraphs:
    if p.text.startswith('Abbreviations:'):
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

# 保存Word文档
output_path = 'Supplementary Table 1.docx'
doc.save(output_path)
