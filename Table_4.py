import pandas as pd
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.enum.section import WD_ORIENTATION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, ns

# Load Brier score data from CSV file
file_path = 'brierscore95.csv'
brier_data = pd.read_csv(file_path)

# Define model names and dataset metrics
model_names = ['RSF', 'CoxPH', 'S-SVM', 'XGBSE', 'GBSA', 'DeepSurv']
metrics = {
    'Train': ['Train integrated', 'Train at 12 months', 'Train at 36 months', 'Train at 60 months'],
    'Test': ['Test integrated', 'Test at 12 months', 'Test at 36 months', 'Test at 60 months'],
    'External': ['External integrated', 'External at 12 months', 'External at 36 months', 'External at 60 months']
}

# Create Word document
doc = Document()

# Set landscape orientation
section = doc.sections[-1]
section.orientation = WD_ORIENTATION.LANDSCAPE
section.page_width, section.page_height = section.page_height, section.page_width

# Add title
title_paragraph = doc.add_paragraph()
run = title_paragraph.add_run('Table 4. ')
run.bold = True  # Bold the "Table 2"
run.font.name = 'Times New Roman'
run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
run = title_paragraph.add_run('Brier score values at 12, 36, and 60 months across different models in three cohorts.')
run.font.name = 'Times New Roman'
run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

# Create table: rows for models and datasets, 5 columns (Model name, Brier Score, 12 months, 36 months, 60 months)
num_rows = len(model_names) * 3 + 1  # Each model has 3 datasets + 1 header row
num_cols = 5
table = doc.add_table(rows=1, cols=num_cols)

# Adjust column width
table.columns[0].width = Inches(2)  # Model name column
for i in range(1, num_cols):
    table.columns[i].width = Inches(2.1)  # Brier score columns

# Function to set table font
def set_table_font(table):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(11)
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

# Create header row
header_cells = table.rows[0].cells
header_cells[0].text = 'Model'
header_cells[1].text = ('Integrated brier score\n(95% CI)')
header_cells[2].text = '12 Months (95% CI)'
header_cells[3].text = '36 Months (95% CI)'
header_cells[4].text = '60 Months (95% CI)'

# Set header row formatting
for i in range(num_cols):
    paragraph = header_cells[i].paragraphs[0]
    run = paragraph.runs[0]
    run.bold = True
    if i == 0:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT  # Left-align first column
    else:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Center-align other columns

# Fill in data: Train, Test, External
for dataset in ['Train', 'Test', 'External']:
    # Add a row for dataset name
    row_cells = table.add_row().cells
    if dataset == 'Train':
        row_cells[0].text = 'Training cohort'
    elif dataset == 'Test':
        row_cells[0].text = 'Test cohort'
    else:
        row_cells[0].text = 'External validation cohort'

    # Merge all columns for the dataset name row
    for i in range(1, num_cols):
        row_cells[i].merge(row_cells[0])

    # Bold and align dataset name
    paragraph = row_cells[0].paragraphs[0]
    run = paragraph.runs[0]
    run.bold = True
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Add rows for each model
    for model in model_names:
        row_cells = table.add_row().cells
        row_cells[0].text = f'{model}'

        # Align "Model" column left
        paragraph = row_cells[0].paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = paragraph.runs[0]
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)

        # Fill other columns with Brier score values
        for i, metric in enumerate(metrics[dataset], start=1):
            value = brier_data.loc[brier_data['Unnamed: 0'] == metric, model].values[0]
            row_cells[i].text = value

        # Align other columns center
        for i in range(1, num_cols):
            row_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

# Apply font settings to the entire table
set_table_font(table)

# Function to set header bottom border (1pt = size 8)
def set_header_bottom_border(table):
    for cell in table.rows[0].cells:
        cell_borders = OxmlElement('w:tcBorders')
        bottom = OxmlElement('w:bottom')
        bottom.set(ns.qn('w:val'), 'single')
        bottom.set(ns.qn('w:sz'), '8')  # 1pt, size 8
        bottom.set(ns.qn('w:color'), '000000')  # Black border
        cell_borders.append(bottom)
        cell._tc.get_or_add_tcPr().append(cell_borders)

# Apply header bottom border
set_header_bottom_border(table)

# Function to set top and bottom border for the entire table (1.5pt = size 12)
def set_table_top_bottom_border(table):
    tbl = table._tbl
    tblBorders = OxmlElement('w:tblBorders')
    top_border = OxmlElement('w:top')
    top_border.set(ns.qn('w:val'), 'single')
    top_border.set(ns.qn('w:sz'), '12')  # 1.5pt, size 12
    top_border.set(ns.qn('w:color'), '000000')
    tblBorders.append(top_border)
    bottom_border = OxmlElement('w:bottom')
    bottom_border.set(ns.qn('w:val'), 'single')
    bottom_border.set(ns.qn('w:sz'), '12')  # 1.5pt, size 12
    bottom_border.set(ns.qn('w:color'), '000000')
    tblBorders.append(bottom_border)
    tbl.tblPr.append(tblBorders)

# Apply table top and bottom border
set_table_top_bottom_border(table)

# Save Word document
output_path = 'Table 4.docx'
doc.save(output_path)

output_path
