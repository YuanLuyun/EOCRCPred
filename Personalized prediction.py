import pandas as pd
import numpy as np
from sklearn.model_selection import train_test_split
from sksurv.ensemble import RandomSurvivalForest
from sksurv.util import Surv
import matplotlib.pyplot as plt
from sklearn.preprocessing import OrdinalEncoder
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.enum.section import WD_ORIENTATION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, ns
from lifelines import KaplanMeierFitter
from lifelines.plotting import add_at_risk_counts
from lifelines.statistics import logrank_test
# 加载数据并去掉 'Patient ID' 列
data = pd.read_csv('data_encoded7408.csv')
data = data.drop(columns=['Patient_ID'])
data_external = pd.read_csv('external_validation_set.csv')
# 构建生存数据
y = Surv.from_dataframe('Survival_status', 'OS_month', data)
X = data.drop(columns=['Survival_status', 'OS_month'])
y_external = Surv.from_dataframe('Survival_status', 'OS_month', data_external)
X_external = data_external.drop(['OS_month', 'Survival_status'], axis=1)

# 按照 7:3 划分训练集与测试集
X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.2, random_state=42)
# 初始化随机生存森林模型
rsf = RandomSurvivalForest(n_estimators=1625, max_depth=6, min_samples_split=2, min_samples_leaf=4, n_jobs=-1,
                           random_state=42)
rsf.fit(X_train, y_train)

# 计算训练集的风险评分
train_risk_scores = rsf.predict(X_train)

# 将训练集的风险评分与生存数据合并
train_df = X_train.copy()
train_df['Risk_Score'] = train_risk_scores
train_df['OS_month'] = y_train['OS_month']
train_df['Survival_status'] = y_train['Survival_status']

# 计算训练集风险评分的三分位数
quantiles = train_df['Risk_Score'].quantile([0.33, 0.67]).values

def risk_category(score, quantiles):
    if score <= quantiles[0]:
        return 'Low risk'
    elif score <= quantiles[1]:
        return 'Medium risk'
    else:
        return 'High risk'

# 对训练集进行风险分层
train_df['Risk_Category'] = train_df['Risk_Score'].apply(lambda x: risk_category(x, quantiles))
# **计算测试集的风险评分**
test_risk_scores = rsf.predict(X_test)

# 对测试集进行同样的处理
test_df = X_test.copy()
test_df['Risk_Score'] = test_risk_scores
test_df['OS_month'] = y_test['OS_month']
test_df['Survival_status'] = y_test['Survival_status']
test_df['Risk_Category'] = test_df['Risk_Score'].apply(lambda x: risk_category(x, quantiles))
# **计算外部验证集的风险评分**
external_risk_scores = rsf.predict(X_external)
# 对外部验证集进行同样的处理
external_df = X_external.copy()
external_df['Risk_Score'] = external_risk_scores
external_df['OS_month'] = y_external['OS_month']
external_df['Survival_status'] = y_external['Survival_status']
external_df['Risk_Category'] = external_df['Risk_Score'].apply(lambda x: risk_category(x, quantiles))

# 初始化 Kaplan-Meier Fitter
kmf = KaplanMeierFitter()

# 设置绘图布局，3列并排显示
fig, axes = plt.subplots(1, 3, figsize=(18, 7))

# 颜色映射
colors = {'Low risk': 'blue', 'Medium risk': 'green', 'High risk': 'red'}

# 计算 log-rank p 值并在图中显示
def add_logrank_p_value(ax, y, data, group1, group2, colors, y_offset=0):
    mask1 = (data['Risk_Category'] == group1)
    mask2 = (data['Risk_Category'] == group2)
    result = logrank_test(y['OS_month'][mask1], y['OS_month'][mask2],
                          event_observed_A=y['Survival_status'][mask1],
                          event_observed_B=y['Survival_status'][mask2])
    p_value = result.p_value
    ax.text(0.02, 0.25 + y_offset, f'{group1} vs {group2} p = {p_value:.3f}',
            transform=ax.transAxes, fontsize=12, color='black')

# 定义每个数据集的 kmf_fitters
kmf_fitters_train = []
kmf_fitters_test = []
kmf_fitters_external = []

for category in ['Low risk', 'Medium risk', 'High risk']:
    # 训练集
    mask_train = (train_df['Risk_Category'] == category)
    kmf_train = KaplanMeierFitter()
    kmf_train.fit(y_train['OS_month'][mask_train], event_observed=y_train[mask_train]['Survival_status'], label=category)
    kmf_fitters_train.append(kmf_train)

    # 测试集
    mask_test = (test_df['Risk_Category'] == category)
    kmf_test = KaplanMeierFitter()
    kmf_test.fit(y_test['OS_month'][mask_test], event_observed=y_test[mask_test]['Survival_status'], label=category)
    kmf_fitters_test.append(kmf_test)

    # 外部验证集
    mask_external = (external_df['Risk_Category'] == category)
    kmf_external = KaplanMeierFitter()
    kmf_external.fit(y_external['OS_month'][mask_external], event_observed=y_external[mask_external]['Survival_status'], label=category)
    kmf_fitters_external.append(kmf_external)

# 绘制训练集生存曲线并计算 log-rank p 值
ax_train = axes[0]
for kmf_train, category in zip(kmf_fitters_train, ['Low risk', 'Medium risk', 'High risk']):
    kmf_train.plot_survival_function(ax=ax_train, color=colors[category])
ax_train.set_title('Training cohort')
ax_train.set_xlabel('Time (months)')
ax_train.set_ylabel('Survival probability')

# 添加风险表
add_at_risk_counts(*kmf_fitters_train, ax=ax_train)

# 添加 log-rank p 值
add_logrank_p_value(ax_train, y_train, train_df, 'Low risk', 'Medium risk', colors, y_offset=0)
add_logrank_p_value(ax_train, y_train, train_df, 'Medium risk', 'High risk', colors, y_offset=0.05)

# 绘制测试集生存曲线并应用 y 轴范围
ax_test = axes[1]
for kmf_test, category in zip(kmf_fitters_test, ['Low risk', 'Medium risk', 'High risk']):
    kmf_test.plot_survival_function(ax=ax_test, color=colors[category])
ax_test.set_title('Test cohort')
ax_test.set_xlabel('Time (months)')
ax_test.set_ylabel('Survival probability')

# 添加风险表
add_at_risk_counts(*kmf_fitters_test, ax=ax_test)

# 添加 log-rank p 值
add_logrank_p_value(ax_test, y_test, test_df, 'Low risk', 'Medium risk', colors, y_offset=0)
add_logrank_p_value(ax_test, y_test, test_df, 'Medium risk', 'High risk', colors, y_offset=0.05)

# 绘制外部验证集生存曲线
ax_external = axes[2]
for kmf_external, category in zip(kmf_fitters_external, ['Low risk', 'Medium risk', 'High risk']):
    kmf_external.plot_survival_function(ax=ax_external, color=colors[category])
ax_external.set_title('External validation cohort')
ax_external.set_xlabel('Time (months)')
ax_external.set_ylabel('Survival probability')

# 添加风险表
add_at_risk_counts(*kmf_fitters_external, ax=ax_external)

# 添加 log-rank p 值
add_logrank_p_value(ax_external, y_external, external_df, 'Low risk', 'Medium risk', colors, y_offset=0)
add_logrank_p_value(ax_external, y_external, external_df, 'Medium risk', 'High risk', colors, y_offset=0.05)

# 在每个图的左上角添加字母
ax_train.text(-0.15, 1.05, 'A', transform=ax_train.transAxes, fontsize=14, fontweight='bold', va='top')
ax_test.text(-0.15, 1.05, 'B', transform=ax_test.transAxes, fontsize=14, fontweight='bold', va='top')
ax_external.text(-0.15, 1.05, 'C', transform=ax_external.transAxes, fontsize=14, fontweight='bold', va='top')

# 调整布局
plt.tight_layout()
plt.savefig('Fig.8.pdf')
plt.close()

print("三张生存曲线图（训练集、测试集、外部验证集）已保存为 'Fig.8.pdf'")

# 读取患者数据
data_2_patients_origin = pd.read_csv('2 patients data.csv')

# 去掉指定的几列
columns_to_drop = ['Risk Score', 'Risk Stratification', '12-month survival rate', '36-month survival rate', '60-month survival rate']
data_2_patients = data_2_patients_origin.drop(columns=columns_to_drop)

# 年龄分区
def categorize_age(age):
    if age < 35:
        return 'less than 35'
    elif 35 <= age < 45:
        return '35-44'
    elif 45 <= age < 50:
        return '45-49'
    else:
        return '50+'  # 超出预期分类的情况

# 对Age列进行分类
data_2_patients['Age'] = data_2_patients['Age'].apply(categorize_age)

# 有序变量及其顺序
ordered_var_categories = {
    'Age': ['less than 35', '35-44', '45-49'],
    'Grade': ['Well differentiated', 'Moderately differentiated', 'Poorly differentiated', 'Undifferentiated'],
    'TNM Stage': ['0', 'I', 'IIA', 'IIB', 'IIC', 'IIIA', 'IIIB', 'IIIC'],
    'T': ['Tis', 'T1', 'T2', 'T3', 'T4'],
    'N': ['N0', 'N1', 'N2'],
    'CEA': ['negative', 'Borderline', 'positive'],
    'No.of resected LNs': ['Zero', '1 to 3', '4+'],
    'Tumor Deposits': ['Zero', '1 to 2', '3+'],
    'Tumor size': ['less than 5', '5+'],
    'Median household income': ['less than $35,000', '$35,000-$54,999', '$55,000-$74,999', '$75,000+']
}

# 使用 OrdinalEncoder 进行有序变量编码
for var, categories in ordered_var_categories.items():
    if var in data_2_patients.columns:
        encoder = OrdinalEncoder(categories=[categories])
        data_2_patients[var] = encoder.fit_transform(data_2_patients[[var]])

# 无序变量及其基准类别
unordered_var_base_category = {
    'Sex': 'Male',
    'Race': 'White',
    'Primary site': 'Ascending colon',
    'Histology': 'Non-specific adenocarcinoma',
    'Resection type': 'Partial/subtotal colectomy',
    'Surg.Rad.Seq': 'Untreated',
    'Chemotherapy': 'No/Unknown',
    'Systemic.Sur.Seq': 'Untreated',
    'Perineural Invasion': 'No',
    'Marital status': 'Single'
}

# 独热编码（无序变量），手动设定基准类别
for var, base_category in unordered_var_base_category.items():
    if var in data_2_patients.columns:
        # 设置类别顺序，将基准类别放在首位
        categories = [base_category] + [cat for cat in data_2_patients[var].unique() if cat != base_category]
        data_2_patients[var] = pd.Categorical(data_2_patients[var], categories=categories, ordered=False)
        # 进行独热编码，并丢弃基准类别
        data_2_patients = pd.get_dummies(data_2_patients, columns=[var], drop_first=True)

# 保存编码后的数据
data_2_patients.to_csv('2 patients data encoded.csv', index=False)

X_2_patients = data_2_patients.drop(['OS', 'Actual survival status'], axis=1)
# 重新排列 data_2_patients 的特征顺序，并填充缺失值为0，确保与 X_train 的特征完全匹配
X_2_patients = X_2_patients.reindex(columns=X_train.columns, fill_value=0)

# 计算生存概率
survival_functions_2_patients = rsf.predict_survival_function(X_2_patients)
# 提取12、36、60个月的生存概率
time_points = [12, 36, 60]
# 计算生存概率
survival_probabilities = {
    time: [survival_functions_2_patients[i](time) for i in range(len(data_2_patients))] for time in time_points
}

# 显示生存概率
for time, probs in survival_probabilities.items():
    print(f'Survival probabilities at {time} months:', probs)

# 对 data_2_patients 进行预测
data_2_patients['Risk_Score'] = rsf.predict(X_2_patients)
data_2_patients['Risk_Category'] = data_2_patients['Risk_Score'].apply(lambda x: risk_category(x, quantiles))

# 输出风险分层结果
print(data_2_patients[['Risk_Score', 'Risk_Category']])

# 创建 Word 文档
doc = Document()

# 设置横向页面
section = doc.sections[-1]
section.orientation = WD_ORIENTATION.LANDSCAPE
section.page_width, section.page_height = section.page_height, section.page_width

# 添加标题
title_paragraph = doc.add_paragraph()
run = title_paragraph.add_run('Table 5. ')
run.bold = True
run.font.name = 'Times New Roman'
run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
run = title_paragraph.add_run(
    'Clinical characteristics and personalized prediction of postoperative survival risk in two EOCRC patients.')
run.font.name = 'Times New Roman'
run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

# 创建表格：3列 (Parameters, Patient A, Patient B)
num_rows = len(data_2_patients) + 1  # 包含表头
num_cols = 3
table = doc.add_table(rows=1, cols=num_cols)

# 调整列宽
for col in table.columns:
    col.width = Inches(2.5)


# 设置表格字体
def set_table_font(table):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(11)
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')


# 添加表头
header_cells = table.rows[0].cells
header_cells[0].text = 'Characteristics'
header_cells[1].text = 'Patient A'
header_cells[2].text = 'Patient B'

# 设置表头字体加粗并居中
for i in range(num_cols):
    paragraph = header_cells[i].paragraphs[0]
    run = paragraph.runs[0]
    run.bold = True
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 确保表头第一列居左对齐
header_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

# 获取第一个患者作为 Patient A，第二个患者作为 Patient B
patient_a_info = data_2_patients_origin.iloc[0]  # 第一位患者
patient_b_info = data_2_patients_origin.iloc[1]  # 第二位患者

# 定义替换字典
replace_dict = {
    'less than $35,000': '＜$35,000',
    '4+': '≥4',
    '3+': '≥3',
    '5+': '≥5',
    'less than 5': '＜5',
    'Zero': '0'
}

# 填充数据，并根据字典进行字符串替换
for column in data_2_patients_origin.columns:  # 遍历参数列
    row_cells = table.add_row().cells

    # 使用字典进行替换
    patient_a_value = str(patient_a_info[column])
    patient_b_value = str(patient_b_info[column])

    row_cells[0].text = column  # 参数名称

    # 如果字典中有对应的值，进行替换，否则保持原值
    row_cells[1].text = replace_dict.get(patient_a_value, patient_a_value)  # Patient A 信息
    row_cells[2].text = replace_dict.get(patient_b_value, patient_b_value)  # Patient B 信息

    # 设置对齐方式
    row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT  # 第一列居左对齐
    row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER  # 第二列居中对齐
    row_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER  # 第三列居中对齐

# 设置表格字体
set_table_font(table)


# 添加表头下方的1磅线
def set_table_border(table):
    # 设置表头上方的线（第一条线）
    tbl = table._tbl
    top_border = OxmlElement('w:top')
    top_border.set(ns.qn('w:val'), 'single')
    top_border.set(ns.qn('w:sz'), '12')  # 1.5pt, size 12
    top_border.set(ns.qn('w:color'), '000000')
    tblBorders = OxmlElement('w:tblBorders')
    tblBorders.append(top_border)

    # 设置表格底部的线（第三条线）
    bottom_border = OxmlElement('w:bottom')
    bottom_border.set(ns.qn('w:val'), 'single')
    bottom_border.set(ns.qn('w:sz'), '12')  # 1.5pt, size 12
    bottom_border.set(ns.qn('w:color'), '000000')
    tblBorders.append(bottom_border)

    tbl.tblPr.append(tblBorders)

# 设置表头底部的边框（第二条线）
def set_header_bottom_border(table):
    header_cells = table.rows[0].cells
    for cell in header_cells:
        tcBorders = OxmlElement('w:tcBorders')
        bottom_border = OxmlElement('w:bottom')
        bottom_border.set(ns.qn('w:val'), 'single')
        bottom_border.set(ns.qn('w:sz'), '8')  # 1pt, size 8
        bottom_border.set(ns.qn('w:color'), '000000')
        tcBorders.append(bottom_border)
        cell._tc.get_or_add_tcPr().append(tcBorders)

def set_row_top_border(table, row_index):
    # 获取目标行
    row_cells = table.rows[row_index].cells
    for cell in row_cells:
        tcBorders = OxmlElement('w:tcBorders')
        top_border = OxmlElement('w:top')  # 设置顶部边框
        top_border.set(ns.qn('w:val'), 'single')
        top_border.set(ns.qn('w:sz'), '8')  # 1pt, size 8
        top_border.set(ns.qn('w:color'), '000000')
        tcBorders.append(top_border)
        cell._tc.get_or_add_tcPr().append(tcBorders)

# 使用此函数为倒数第七行上方添加边框
set_row_top_border(table, len(table.rows) - 7)

# 应用三线表边框
set_table_border(table)
set_header_bottom_border(table)  # 设置表头底部边框
# 保存表格之前，在表格下方添加注释
paragraph = doc.add_paragraph()
run = paragraph.add_run('Abbreviations:')
run.bold = True  # 加粗 "Abbreviations"
run.font.name = 'Times New Roman'  # 设置字体为 Times New Roman
run.font.size = Pt(11)
run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

# 添加缩写解释
abbreviation_text = """
CEA: Carcinoembryonic Antigen; No. of resected LNs: Number of Resected Lymph Nodes; OS: Overall Survival; Surg.Rad.Seq: Surgery Radiation Sequence; Systemic.Sur.Seq: Systemic Surgery Sequence
"""

paragraph.add_run(abbreviation_text).font.size = Pt(11)

# 设置注释字体和格式
for p in doc.paragraphs:
    if p.text.startswith('Abbreviations:'):
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

# 保存 Word 文档
doc.save('Table 5.docx')

# 假设 rsf 是已经训练好的随机生存森林模型，X_2_patients 是包含两个患者特征的数据

# 使用 predict_survival_function 获取生存曲线
survival_functions_2_patients = rsf.predict_survival_function(X_2_patients)

# 生成并排的累积风险曲线
fig, axs = plt.subplots(1, 2, figsize=(14, 6))
time_points = np.linspace(0, 120, num=121)

# 初始化变量存储所有患者的最大累计风险
max_cumulative_hazard = 0

# 绘制两位患者的累积风险曲线，并同时找到最大风险值
for i, survival_function in enumerate(survival_functions_2_patients):
    # 计算每位患者的累计风险曲线
    cumulative_hazard = [1 - survival_function(t) for t in time_points]

    # 更新最大累计风险值
    max_cumulative_hazard = max(max_cumulative_hazard, max(cumulative_hazard))

    # 绘制累计风险曲线
    axs[i].step(time_points, cumulative_hazard, where='post')

    # 设置标题、轴标签和网格
    axs[i].set_title(f'Patient {chr(65 + i)}')  # 使用 A 和 B
    axs[i].set_xlabel('Time (months)')
    axs[i].set_ylabel('Cumulative hazard')
    axs[i].grid()

# 设置y轴范围，以所有患者的最大累积风险为准
for ax in axs:
    ax.set_ylim(0, max_cumulative_hazard)  # 设置相同的y轴范围

# 调整布局并保存图像
plt.tight_layout()
plt.savefig('Fig.9.pdf')
plt.close()

# 输出提示
print("累积风险曲线已保存为 Fig.9.pdf")


